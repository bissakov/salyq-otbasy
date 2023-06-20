import time
from time import sleep
import datetime
import json
import os
import pdfkit
import psutil
import pywinauto
import requests
import shutil
import urllib3
import win32com.client
from datetime import timedelta
from docx import Document
from docx.shared import Inches
from mail import send_email
from os import listdir, makedirs
from os.path import join, exists
from selenium import webdriver
from selenium.common import TimeoutException, NoSuchWindowException
from selenium.webdriver import Keys, ActionChains
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.webdriver import WebDriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from telegram_send import send_message
from typing import Tuple, Dict, List, Any, Optional
from webdriver_manager.chrome import ChromeDriverManager
from dotenv import load_dotenv
from urllib.parse import urljoin


load_dotenv()
urllib3.disable_warnings()

today: str = datetime.datetime.now().strftime('%d.%m.%Y')

base_url: str = os.getenv('BASE_URL')
index_page_url: str = urljoin(base_url, 'knp/main/index')
personal_cabinet_url: str = urljoin(base_url, 'knp/personal-cabinet/')
logout_url: str = urljoin(base_url, 'sonoweb/content/exit.faces')
notification_url: str = urljoin(base_url, 'knp/notifications/esutd/?id=')
tax_statements_url: str = urljoin(base_url, 'knp/declarations/registry/')

login_button: Tuple[str, str] = (By.CSS_SELECTOR, 'button.enter-by-cert-button')
auth_key_input: Tuple[str, str] = (By.CSS_SELECTOR, 'input.custom-file-input')
password_input: Tuple[str, str] = (By.CSS_SELECTOR, 'input.listBox.form-control')
confirm_password_button: Tuple[str, str] = (By.CSS_SELECTOR, 'div.input-group-append > button')
confirm_login_button: Tuple[str, str] = (By.CSS_SELECTOR, 'button.btn.btn-primary')
user_info: Tuple[str, str] = (By.CSS_SELECTOR, 'div.userInfo')
page_title: Tuple[str, str] = (By.CSS_SELECTOR, 'h1.pageTitle')
news_date: Tuple[str, str] = (By.CSS_SELECTOR, 'div.news-item__date')
calendar: Tuple[str, str] = (By.CSS_SELECTOR, 'div.b-calendar')
menu: Tuple[str, str] = (By.CSS_SELECTOR, 'table.menuTable')
from_date_button: Tuple[str, str] = (By.CSS_SELECTOR, 'button#__BVID__44')
to_date_button: Tuple[str, str] = (By.CSS_SELECTOR, 'button#__BVID__49')
today_date_dropdown: Tuple[str, str] = (By.CSS_SELECTOR, f'div[aria-label="{today} (Сегодня)"]')
submit_button: Tuple[str, str] = (By.CSS_SELECTOR, 'button[type="submit""]')
loader: Tuple[str, str] = (By.CSS_SELECTOR, 'div.thin-loader')
notification_content: Tuple[str, str] = (By.CSS_SELECTOR, '#notification-content')

screen_local_folder_path: str = fr'Screen/{today}'
screen_fserver_folder_path: str = fr'\\fserver\ДБУИО_Новая\ДБУИО_Общая папка\ДБУ_Информация УНУ\АКТ СВЕРКИ, ЛЧС\scrin-2023'
notification_folder_path: str = fr'C:\Users\robot.ad\PycharmProjects\Salyk\Notifications\{today}'
base_path: str = r'\\fserver\ДБУИО_Новая\ДБУИО_Общая папка\ДБУ_Информация УНУ\Криптоключи действующие\Робот'
pdf_save_path: str = r'\\fserver\ДБУИО_Новая\ДБУИО_Общая папка\ДБУ_Информация УНУ\АКТ СВЕРКИ, ЛЧС\01_Анализ лицевых счетов\СПРАВКА об отсутсв зад-ти\2023'


def login(driver: WebDriver, wait: WebDriverWait, auth_key_path: str, password: str) -> None:
    job_done = False
    while not job_done:
        try:
            driver.get(index_page_url)
            wait.until(EC.element_to_be_clickable(login_button)).click()
            sleep(1)
            wait.until(EC.presence_of_element_located(auth_key_input)).send_keys(auth_key_path)
            sleep(1)
            wait.until(EC.element_to_be_clickable(password_input)).send_keys(password)
            sleep(1)
            wait.until(EC.element_to_be_clickable(confirm_password_button)).click()
            sleep(1)
            wait.until(EC.element_to_be_clickable(confirm_login_button)).click()

            sleep(10)

            if 'notification' in driver.current_url:
                driver.find_element(By.CSS_SELECTOR, 'input[type="checkbox"]').click()
                sleep(1)
                driver.find_element(By.CSS_SELECTOR, 'button').click()

            while not job_done:
                try:
                    wait.until(EC.presence_of_element_located(user_info))
                    wait.until(EC.presence_of_element_located(page_title))
                    wait.until(EC.presence_of_element_located(news_date))
                    wait.until(EC.presence_of_element_located(calendar))
                    wait.until(EC.presence_of_element_located(menu))
                    job_done = True
                except (TimeoutException, NoSuchWindowException):
                    if 'personal-cabinet' in driver.current_url:
                        driver.refresh()
                    else:
                        break
        except TimeoutException:
            sleep(5)
            continue


def get_headers(driver: WebDriver) -> Dict[str, str]:
    cookie_session = driver.get_cookie('cookiesession1')['value']
    ns = driver.get_cookie('NS')['value']
    nsiv = driver.get_cookie('NSIV')['value']
    return {
        'Accept': '*/*',
        'Accept-Language': 'en-US,en;q=0.9',
        'Connection': 'keep-alive',
        'Content-Type': 'application/json;charset=utf-8',
        'Cookie': f'cookiesession1={cookie_session}; NS={ns}; NSIV={nsiv}',
        'Origin': base_url,
        'Referer': urljoin(base_url, 'knp/notifications/registry/'),
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-origin',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
                      'AppleWebKit/537.36 (KHTML, like Gecko) Chrome/113.0.0.0 Safari/537.36',
    }


def get_latest_working_day() -> str:
    today_date = datetime.datetime.strptime(today, '%d.%m.%Y').date()
    prev_date = today_date - timedelta(days=1)

    if prev_date.weekday() >= 5:
        prev_date -= datetime.timedelta(days=prev_date.weekday() - 4)

    return prev_date.strftime('%d.%m.%Y')


def get_notifications(driver: WebDriver) -> Optional[List[Any]]:
    url = urljoin(base_url, 'notifications/registry/tp/list')

    payload = {
        'receiveDate1': get_latest_working_day(),
        'receiveDate2': today,
        'readDate1': None,
        'readDate2': None,
        'notificationNumber': None,
        'page': 1,
        'pageSize': 10,
        'pageSortBy': 'receiveDate',
        'pageSortAsc': False
    }
    headers = get_headers(driver=driver)
    response = None

    while True:
        try:
            response = requests.request('POST', url, json=payload, headers=headers, verify=False)
        except requests.exceptions.ConnectionError:
            sleep(60)
        break
    try:
        return response.json()
    except (json.decoder.JSONDecodeError, AttributeError):
        return None


def save_screen_doc(branch_mappings: Dict[str, str]) -> None:
    screen_doc_path = join(screen_local_folder_path, f'{today}.docx')
    doc = Document()
    doc.add_heading('Снимки экрана Salyq', 0)
    for branch, branch_name in branch_mappings.items():
        doc.add_paragraph(branch_name)
        doc.add_picture(join(screen_local_folder_path, f'{branch}.png'), width=Inches(7))
    doc.save(screen_doc_path)
    shutil.copyfile(screen_doc_path, join(screen_fserver_folder_path, f'{today}.docx'))


def get_current_process_pid(proc_name: str) -> int or None:
    return next((p.pid for p in psutil.process_iter() if proc_name in p.name()), None)


def kill_all_processes(proc_name: str) -> None:
    for proc in psutil.process_iter():
        if proc_name in proc.name():
            process = psutil.Process(proc.pid)
            try:
                process.terminate()
            except psutil.AccessDenied:
                continue


def save_notification_doc(prefix: str) -> None:
    job_done = False
    while not job_done:
        try:
            kill_all_processes(proc_name='WINWORD.EXE')
            notification_doc_file = join(notification_folder_path, f'уведомление_{prefix}.docx')
            if exists(notification_doc_file):
                os.unlink(notification_doc_file)

            word = win32com.client.Dispatch("Word.Application")
            word.visible = True
            word.DisplayAlerts = False

            if not exists(notification_doc_file):
                with open(notification_doc_file, 'w'):
                    pass
            word.Documents.Open(notification_doc_file)
            sleep(5)
            word_pid = get_current_process_pid(proc_name='WINWORD.EXE')
            app = pywinauto.Application(backend='uia').connect(process=word_pid)
            for win in app.windows():
                win_text = win.window_text()
                if not win_text:
                    continue
                window = app.window(title=win_text)
                window['Закрыть'].click()
            doc = word.ActiveDocument
            app.top_window().set_focus()

            app.top_window().type_keys('{VK_CONTROL down}v{VK_CONTROL up}')
            doc.Close(True)
            word.Quit()
            job_done = True
        except AttributeError:
            continue


def save_notification(notification: Dict[str, any], driver: WebDriver, wait: WebDriverWait, branch: str) -> bool:
    if notification['descriptionRu'] == 'низкая':
        return False
    notification_id = notification['id']
    receive_date = notification['receiveDate'].replace(':', '')
    job_done = False
    retry_count = 0
    while not job_done:
        if retry_count == 3:
            return False
        try:
            driver.get(f'{notification_url}{notification_id}')
            wait.until(EC.presence_of_element_located(notification_content))
            job_done = True
        except TimeoutException:
            sleep(5)
            retry_count += 1
            continue
    driver.execute_script('''
        const notification = document.querySelector('#notification-content > div > div > div');
        document.querySelector('body > div').remove();
        document.body.append(notification);
    ''')
    ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').send_keys('c').perform()
    save_notification_doc(prefix=f'{branch}_{receive_date}')
    return True


def send_tax_request(session: requests.Session, headers: Dict[str, str]):
    url = urljoin(base_url, 'declaration/debt/send')
    payload = {
        'dateRequest': today,
        'refGoal': '0xffff00000019',
        'refReceiver': '0xffff00000011',
        'taxOrgCode': '6007'
    }
    response = session.request('POST', url, json=payload, headers=headers, verify=False)
    response.raise_for_status()


def save_pdf_statement(doc_info: Dict, session: requests.Session, headers: Dict[str, str]) -> None:
    url = urljoin(base_url, f'{doc_info["actions"][0]["target"]}')
    response = session.request('GET', url, headers=headers, verify=False)

    prefix = len(listdir(pdf_save_path)) + 1
    with open(join(pdf_save_path, f'{prefix}_справка_{today}.pdf'), mode='wb') as pdf_file:
        pdf_file.write(response.content)


def get_tax_statement(session: requests.Session, headers: Dict[str, str]):
    while True:
        url = urljoin(base_url, 'declarations/registry/tp/allByDates')
        querystring = {'from': today, 'to': today}
        response = session.request('GET', url, headers=headers, params=querystring, verify=False)
        docs_infos: List[Dict[str, str]] = response.json()
        print(docs_infos)
        if not docs_infos:
            continue
        if len(docs_infos[0]['actions']) != 0:
            save_pdf_statement(doc_info=docs_infos[0], session=session, headers=headers)
            break


def run_salyq() -> None:
    # send_message('Старт процесса Salyq')

    with open(file='branch_mapping.json', mode='r', encoding='utf-8') as f:
        branch_mappings: Dict[str, str] = json.load(f)
    if not exists(screen_local_folder_path):
        makedirs(screen_local_folder_path)
    if not exists(notification_folder_path):
        makedirs(notification_folder_path)

    service = Service(executable_path=ChromeDriverManager().install())
    options = webdriver.ChromeOptions()
    options.add_argument('--start-maximized')
    driver = webdriver.Chrome(service=service, options=options)
    wait = WebDriverWait(driver, 10)

    branch_mappings = {key: val for key, val in branch_mappings.items() if int(key) >= 18}

    with driver:
        for branch, branch_name in branch_mappings.items():

            print(branch)

            password_folder = join(base_path, branch)
            password = listdir(password_folder)[0]
            auth_key_path = join(password_folder, password, listdir(join(password_folder, password))[0])

            login(driver=driver, wait=wait, auth_key_path=auth_key_path, password=password)
            driver.save_screenshot(fr'{screen_local_folder_path}/{branch}.png')
            send_message(f'Сохранен скриншот по филиалу {branch}')

            notifications = get_notifications(driver=driver)

            if notifications:
                print(f'Notifications: {len(notifications)}')
                for notification in notifications:
                    if save_notification(notification=notification, driver=driver, wait=wait, branch=branch):
                        send_message(f'🟢Есть уведомление по филиалу {branch}🟢')
                    else:
                        if notification['descriptionRu'] != 'низкая':
                            url = urljoin(base_url, f'notifications/inis/view/{notification["id"]}')
                            headers = get_headers(driver=driver)
                            response = requests.request('GET', url, headers=headers, verify=False)
                            response.raise_for_status()
                            receive_date = notification['receiveDate'].replace(':', '')
                            pdf_name = join(notification_folder_path, f'уведомление_{branch}_{receive_date}.pdf')
                            pdfkit.from_string(response.text.strip(), r'C:\temp.pdf', options={"encoding": "UTF-8"})
                            shutil.copyfile(r'C:\temp.pdf', pdf_name)
                            os.unlink(r'C:\temp.pdf')
            else:
                print('No notifications')
                send_message(f'Нет уведомлений по {branch}')

            driver.get(logout_url)
            sleep(2)

    send_email(folder_path=notification_folder_path)
    save_screen_doc(branch_mappings=branch_mappings)

    if datetime.datetime.now().weekday() == 4:
        service = Service(executable_path=ChromeDriverManager().install())
        options = webdriver.ChromeOptions()
        options.add_argument('--start-maximized')
        driver = webdriver.Chrome(service=service, options=options)
        wait = WebDriverWait(driver, 10)
        with driver:
            branch = '18'
            print(branch)

            password_folder = join(base_path, branch)
            password = listdir(password_folder)[0]
            auth_key_path = join(password_folder, password, listdir(join(password_folder, password))[0])

            login(driver=driver, wait=wait, auth_key_path=auth_key_path, password=password)

            headers = get_headers(driver=driver)
            with requests.Session() as session:
                send_message('Отправление запроса на получение сведений об отсутствии задолженности')
                send_tax_request(session=session, headers=headers)
                sleep(5)
                send_message('Ожидание обработки и получение PDF документа')
                get_tax_statement(session=session, headers=headers)
                send_message('Справка успешно сохранилась')
    send_message('Конец процесса Salyq')


def wait_until(target_hour: int) -> None:
    while True:
        current_hour = time.localtime().tm_hour
        if current_hour == target_hour:
            break
        print('waiting')
        sleep(300)


if __name__ == '__main__':
    try:
        wait_until(target_hour=9)
        run_salyq()
    except Exception as e:
        send_message(message=str(e), is_error=True)
        raise e
