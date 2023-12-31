from time import sleep

import datetime
import logging
import pdfkit
import requests
import shutil
from agent_initialization import *
from constants import *
from datetime import timedelta
from docx import Document
from docx.shared import Inches
from mail import send_email
from os import listdir, makedirs
from os.path import join
from selenium import webdriver
from selenium.common import TimeoutException, NoSuchWindowException
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from telegram_send import send_message
from typing import Dict, List, Any, Optional
from webdriver_manager.chrome import ChromeDriverManager


def login(driver: WebDriver, wait: WebDriverWait, auth_key_path: str, password: str) -> None:
    job_done = False
    while not job_done:
        try:
            driver.get(INDEX_PAGE_URL)
            logging.info(f'Index page {INDEX_PAGE_URL} opened')
            wait.until(EC.element_to_be_clickable(LOGIN_BUTTON)).click()
            logging.info(f'Login button clicked')
            sleep(1)
            wait.until(EC.presence_of_element_located(AUTH_KEY_INPUT)).send_keys(auth_key_path)
            logging.info(f'Auth key path {auth_key_path} entered')
            sleep(1)
            wait.until(EC.element_to_be_clickable(PASSWORD_INPUT)).send_keys(password)
            logging.info(f'Password {password} entered')
            sleep(1)
            wait.until(EC.element_to_be_clickable(CONFIRM_PASSWORD_BUTTON)).click()
            logging.info('Confirm password button clicked')
            sleep(1)
            wait.until(EC.element_to_be_clickable(CONFIRM_LOGIN_BUTTON)).click()
            logging.info('Confirm login button clicked')

            sleep(10)

            if 'notification' in driver.current_url:
                logging.info(f'Redirected to notification page {driver.current_url}')
                driver.find_element(By.CSS_SELECTOR, 'input[type="checkbox"]').click()
                logging.info('Checkbox clicked')
                sleep(1)
                driver.find_element(By.CSS_SELECTOR, 'button').click()
                logging.info('Confirm button clicked')

            while not job_done:
                try:
                    wait.until(EC.presence_of_element_located(USER_INFO))
                    wait.until(EC.presence_of_element_located(PAGE_TITLE))
                    wait.until(EC.presence_of_element_located(NEWS_DATE))
                    wait.until(EC.presence_of_element_located(CALENDAR))
                    wait.until(EC.presence_of_element_located(MENU))
                    logging.info('Personal cabinet page successfully loaded')
                    job_done = True
                except (TimeoutException, NoSuchWindowException) as error:
                    logging.info('Elements from personal cabinet page not loaded')
                    if 'personal-cabinet' in driver.current_url:
                        logging.info(f'Refreshing {driver.current_url} because of error {error}')
                        driver.refresh()
                    else:
                        logging.info(f'Starting login process over because of error {error}')
                        break
        except TimeoutException as error:
            logging.info(f'Starting login process over because of error {error}')
            sleep(5)
            continue


def get_headers(driver: WebDriver) -> Dict[str, str]:
    ns = driver.get_cookie('NS')['value']
    nsiv = driver.get_cookie('NSIV')['value']
    return {
        'Accept': '*/*',
        'Accept-Language': 'en-US,en;q=0.9',
        'Connection': 'keep-alive',
        'Content-Type': 'application/json;charset=utf-8',
        'Cookie': f'NS={ns}; NSIV={nsiv}',
        'Origin': BASE_URL,
        'Referer': TAX_STATEMENTS_URL,
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-origin',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
                      'AppleWebKit/537.36 (KHTML, like Gecko) Chrome/113.0.0.0 Safari/537.36',
    }


def get_latest_working_day(today: str) -> str:
    today_date = datetime.datetime.strptime(today, '%d.%m.%Y').date()
    prev_date = today_date - timedelta(days=1)

    if prev_date.weekday() >= 5:
        prev_date -= datetime.timedelta(days=prev_date.weekday() - 4)

    return prev_date.strftime('%d.%m.%Y')


def get_notifications(driver: WebDriver, today: str) -> Optional[List[Any]]:
    logging.info('Checking notifications')

    url = urljoin(BASE_URL, 'notifications/registry/tp/list')
    logging.info(f'Notifications url {url}')

    payload = {
        'receiveDate1': get_latest_working_day(today=today),
        'receiveDate2': today,
        'readDate1': None,
        'readDate2': None,
        'notificationNumber': None,
        'page': 1,
        'pageSize': 10,
        'pageSortBy': 'receiveDate',
        'pageSortAsc': False
    }
    logging.info(f'Notifications payload {payload}')

    headers = get_headers(driver=driver)
    logging.info(f'Notifications headers {headers}')

    response = None

    while True:
        try:
            response = requests.request('POST', url, json=payload, headers=headers, verify=False)
            logging.info(f'Notifications response {response}')
        except requests.exceptions.ConnectionError as error:
            logging.info(f'Sleeping for 60 seconds. Notifications connection error {error}')
            sleep(60)
        break
    try:
        logging.info(f'Notifications response json {response.json()}')
        return response.json()
    except (json.decoder.JSONDecodeError, AttributeError) as error:
        logging.info(f'Notifications response json error {error}')
        return None


def save_screen_doc(today: str, branch_mappings: Dict[str, str]) -> None:
    screen_local_folder_path = join(SCREEN_LOCAL_FOLDER_PATH, today)
    screen_doc_path = join(screen_local_folder_path, f'{today}.docx')
    doc = Document()
    doc.add_heading('Снимки экрана Salyk', 0)
    for branch, branch_name in branch_mappings.items():
        doc.add_paragraph(branch_name)
        doc.add_picture(join(screen_local_folder_path, f'{branch}.png'), width=Inches(7))
    doc.save(screen_doc_path)
    file_name = f'{today}.docx'
    current_day = datetime.datetime.now().strftime('%d.%m.%Y')
    if today != current_day:
        file_name = f'{today}_({current_day}).docx'
    shutil.copyfile(screen_doc_path, join(SCREEN_FSERVER_FOLDER_PATH, file_name))


def save_notification(today: str, notification: Dict[str, any],
                      driver: WebDriver, wait: WebDriverWait, branch: str) -> bool:
    if notification['descriptionRu'] == 'низкая':
        logging.info('Invalid type of notification ("низкая")')
        return False
    notification_id = notification['id']
    logging.info(f'Notification id {notification_id}')
    receive_date = notification['receiveDate'].replace(':', '')
    logging.info(f'Notification receive date {receive_date}')
    job_done = False
    retry_count = 0
    while not job_done:
        if retry_count == 3:
            logging.info(f'Notification could not be opened. Number of tries exceeded {retry_count}')
            return False
        try:
            notification_url = f'{NOTIFICATION_URL}{notification_id}'
            driver.get(notification_url)
            logging.info(f'Notification url {notification_url}')
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '.mainContent')))
            wait.until(EC.text_to_be_present_in_element((By.TAG_NAME, 'body'), 'Уважаемый работодатель!'))
            logging.info('Notification content is present')
            job_done = True
        except TimeoutException:
            sleep(5)
            retry_count += 1
            logging.info(f'Notification content is not present. Retrying: {retry_count}')
            continue
    driver.execute_script('''
        const notificationContentElement = document.querySelector('#notification-content');
        const textNormalElement = document.querySelector('.textNormal');

        let notification;

        if (notificationContentElement !== null) {
          notification = document.querySelector('#notification-content > div > div > div')
        } else if (textNormalElement !== null) {
          notification = document.querySelector('.textNormal');
        }  

        document.querySelector('body > div').remove();
        document.body.append(notification);

        const scripts = document.querySelectorAll('script, noscript');
        scripts.forEach(script => script.remove());
    ''')

    notification_folder_path = join(NOTIFICATION_FOLDER_PATH, today)
    pdf_name = join(notification_folder_path, f'уведомление_{branch}_{receive_date}.pdf')
    logging.info(f'pdf_name: {pdf_name}')

    source = driver.execute_script('return document.body.outerHTML;')
    pdfkit.from_string(source, r'C:\temp.pdf', options={'encoding': 'UTF-8'})
    shutil.copyfile(r'C:\temp.pdf', pdf_name)
    return True


def save_notification_risk(today: str, notification: Dict[str, any],
                           driver: WebDriver, wait: WebDriverWait, branch: str) -> bool:
    if notification['descriptionRu'] == 'низкая':
        logging.info('Invalid type of notification ("низкая")')
        return False
    notification_id = notification['id']
    logging.info(f'Notification id {notification_id}')
    receive_date = notification['receiveDate'].replace(':', '')
    logging.info(f'Notification receive date {receive_date}')
    job_done = False
    retry_count = 0
    while not job_done:
        if retry_count == 3:
            logging.info(f'Notification could not be opened. Number of tries exceeded {retry_count}')
            return False
        try:
            notification_url = f'https://cabinet.salyk.kz/knp/notifications/risk/?id={notification_id}'
            driver.get(notification_url)
            logging.info(f'Notification url {notification_url}')
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '#risk-content')))
            logging.info('Notification content is present')
            job_done = True
        except TimeoutException:
            sleep(5)
            retry_count += 1
            logging.info(f'Notification content is not present. Retrying: {retry_count}')
            continue
    driver.execute_script('''
        document.querySelector('.print-button').remove();
        const notification = document.querySelector('#risk-content');
        document.querySelector('body > div').remove();
        document.body.append(notification);
        
        const scripts = document.querySelectorAll('script, noscript');
        scripts.forEach(script => script.remove());
    ''')

    notification_folder_path = join(NOTIFICATION_FOLDER_PATH, today)
    pdf_name = join(notification_folder_path, f'уведомление_{branch}_{receive_date}.pdf')
    logging.info(f'pdf_name: {pdf_name}')

    source = driver.execute_script('return document.body.outerHTML;')
    pdfkit.from_string(source, r'C:\temp.pdf', options={'encoding': 'UTF-8'})
    shutil.copyfile(r'C:\temp.pdf', pdf_name)
    return True


def send_tax_request(today: str, session: requests.Session, headers: Dict[str, str]):
    url = urljoin(BASE_URL, 'declaration/debt/send')
    payload = {
        'dateRequest': today,
        'refGoal': '0xffff00000019',
        'refReceiver': '0xffff00000011',
        'taxOrgCode': '6007'
    }
    response = session.request('POST', url, json=payload, headers=headers, verify=False)
    response.raise_for_status()


def save_pdf_statement(today: str, doc_info: Dict, session: requests.Session, headers: Dict[str, str]) -> None:
    url = urljoin(BASE_URL, f'{doc_info["actions"][0]["target"]}')
    response = session.request('GET', url, headers=headers, verify=False)

    prefix = len(listdir(PDF_SAVE_PATH)) + 1
    with open(join(PDF_SAVE_PATH, f'{prefix}_справка_{today}.pdf'), mode='wb') as pdf_file:
        pdf_file.write(response.content)


def get_tax_statement(today: str, session: requests.Session, headers: Dict[str, str]):
    while True:
        url = urljoin(BASE_URL, 'declarations/registry/tp/allByDates')
        querystring = {'from': today, 'to': today}
        response = session.request('GET', url, headers=headers, params=querystring, verify=False)
        docs_infos: List[Dict[str, str]] = response.json()
        print(docs_infos)
        if not docs_infos:
            continue
        if len(docs_infos[0]['actions']) != 0:
            save_pdf_statement(today=today, doc_info=docs_infos[0], session=session, headers=headers)
            break


def run_salyk(today: str) -> None:
    send_message(f'Старт процесса Salyk за {today}')

    with open(file=r'C:\Users\robot.ad\PycharmProjects\Salyk\branch_mapping.json', mode='r', encoding='utf-8') as f:
        branch_mappings: Dict[str, str] = json.load(f)

    logging.info(f'branch_mappings: {branch_mappings}')

    screen_local_folder_path = join(SCREEN_LOCAL_FOLDER_PATH, today)
    notification_folder_path = join(NOTIFICATION_FOLDER_PATH, today)

    logging.info(f'screen_local_folder_path: {screen_local_folder_path}')
    logging.info(f'notification_folder_path: {notification_folder_path}')

    makedirs(screen_local_folder_path, exist_ok=True)
    makedirs(notification_folder_path, exist_ok=True)

    logging.info('Созданы папки для скриншотов и уведомлений')

    try:
        service = Service(executable_path=ChromeDriverManager().install())
    except Exception as e:
        service = Service(executable_path=ChromeDriverManager(version='114.0.5735.16').install())
    options = webdriver.ChromeOptions()
    options.add_argument('--start-maximized')
    driver = webdriver.Chrome(service=service, options=options)
    wait = WebDriverWait(driver, 10)

    logging.info('Chrome launched')

    # branch_mappings = {key: val for key, val in branch_mappings.items() if int(key) >= 18}

    with driver:
        for branch, branch_name in branch_mappings.items():
            logging.info(f'Working on a branch: {branch}')

            password_folder = join(BASE_PATH, branch)
            logging.info(f'password_folder: {password_folder}')
            password = listdir(password_folder)[0]
            logging.info(f'password: {password}')
            auth_key_path = join(password_folder, password, listdir(join(password_folder, password))[0])
            logging.info(f'auth_key_path: {auth_key_path}')

            logging.info(f'Logging into Salyk with a brach {branch}')
            login(driver=driver, wait=wait, auth_key_path=auth_key_path, password=password)

            screenshot_path = fr'{screen_local_folder_path}/{branch}.png'
            driver.save_screenshot(screenshot_path)
            logging.info(f'Screenshot {screenshot_path} saved for a branch {branch}')
            send_message(f'Сохранен скриншот по филиалу {branch}')

            notifications = get_notifications(driver=driver, today=today)

            if notifications:
                logging.info(f'Notifications: {len(notifications)} for branch: {branch}')
                for notification in notifications:
                    if save_notification(today=today, notification=notification,
                                         driver=driver, wait=wait, branch=branch):
                        logging.info(f'Notification {branch} saved via manual copy and pasting')
                        send_message(f'Есть уведомление по филиалу {branch}')
                    elif save_notification_risk(today=today, notification=notification,
                                                driver=driver, wait=wait, branch=branch):
                        logging.info(f'Notification {branch} saved via manual copy and pasting')
                        send_message(f'Есть уведомление по филиалу {branch}')
                    else:
                        if notification['descriptionRu'] != 'низкая':
                            send_message(f'Есть уведомление по филиалу {branch}')
                            logging.info(f'Branch {branch} has a notification that could be saved via request')

                            url = urljoin(BASE_URL, f'notifications/inis/view/{notification["id"]}')
                            logging.info(f'Notification url: {url}')

                            headers = get_headers(driver=driver)
                            logging.info(f'Headers: {headers}')

                            response = requests.request('GET', url, headers=headers, verify=False)
                            logging.info(f'Response status code: {response.status_code}')

                            response.raise_for_status()

                            receive_date = notification['receiveDate'].replace(':', '')
                            logging.info(f'receive_date: {receive_date}')

                            pdf_name = join(notification_folder_path, f'уведомление_{branch}_{receive_date}.pdf')
                            logging.info(f'pdf_name: {pdf_name}')

                            if response.text == '':
                                url = urljoin(
                                    BASE_URL,
                                    f'notifications/cc/tp/notification/download/{notification["id"]}'
                                )
                                response = requests.request('GET', url, headers=headers, verify=False)
                                with open(pdf_name, 'wb') as f:
                                    f.write(response.content)
                            else:
                                pdfkit.from_string(response.text.strip(), r'C:\temp.pdf', options={"encoding": "UTF-8"})
                                shutil.copyfile(r'C:\temp.pdf', pdf_name)

                            logging.info(f'PDF {pdf_name} saved for a branch {branch}')

                            os.unlink(r'C:\temp.pdf')
            else:
                logging.info(f'No notifications for a branch {branch}')
                send_message(f'Нет уведомлений по {branch}')

            driver.get(LOGOUT_URL)
            sleep(2)

    attachments = [join(notification_folder_path, file_name) for file_name in listdir(notification_folder_path)]
    if attachments:
        send_message('Отправка уведомлений')
        send_email(attachments=attachments)
        logging.info('Email sent')

    save_screen_doc(today=today, branch_mappings=branch_mappings)
    doc_path = join(SCREEN_FSERVER_FOLDER_PATH, f'{today}.docx')
    send_message(f'Сохранен документ со скринштоами {doc_path}')
    logging.info('Screen doc saved')

    if datetime.datetime.now().weekday() == 4:
        try:
            service = Service(executable_path=ChromeDriverManager().install())
        except Exception as e:
            service = Service(executable_path=ChromeDriverManager(version='114.0.5735.16').install())
        options = webdriver.ChromeOptions()
        options.add_argument('--start-maximized')
        driver = webdriver.Chrome(service=service, options=options)
        wait = WebDriverWait(driver, 10)
        with driver:
            branch = '18'
            print(branch)

            password_folder = join(BASE_PATH, branch)
            password = listdir(password_folder)[0]
            auth_key_path = join(password_folder, password, listdir(join(password_folder, password))[0])

            login(driver=driver, wait=wait, auth_key_path=auth_key_path, password=password)

            headers = get_headers(driver=driver)
            with requests.Session() as session:
                send_message('Отправление запроса на получение сведений об отсутствии задолженности')
                send_tax_request(today=today, session=session, headers=headers)
                sleep(5)
                send_message('Ожидание обработки и получение PDF документа')
                get_tax_statement(today=today, session=session, headers=headers)
                send_message('Справка успешно сохранилась')
    send_message('Конец процесса Salyk')


def run():
    logging.info('Salyk process started')
    try:
        today_date = datetime.datetime.now()
        logging.info(f'Current date: {today_date}')

        logging.info('Starting Salyk process')
        today = datetime.datetime.now().strftime('%d.%m.%Y')
        run_salyk(today=today)
        logging.info('Salyk process finished.')
    except Exception as e:
        error_msg = f'Error occured on robot-12\n' \
                    f'Process: Salyk\n' \
                    f'Error:\n{e}'
        send_message(message=error_msg, is_error=True)
        logging.error(e)
        raise e
